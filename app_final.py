# =============================================================================
# Plantilla institucional UCCuyo (Streamlit) — checklist al portar otra repo:
#   • assets/escudo_uccuyo.png en repo + URL raw pública GitHub (src del escudo en HTML)
#   • Un solo HTML del banner (.ucci-inst-header) con fondo verde también en
#     style="" (no depende de :has() ni del DOM de st.container horizontal).
#   • .streamlit/config.toml [theme] (secondaryBackgroundColor del cargador)
# =============================================================================
import os
from pathlib import Path

import streamlit as st
import pandas as pd
import pdfplumber
import yaml
import io
from docx import Document
from docx.shared import Pt
from datetime import datetime
from openpyxl import Workbook

_APP_DIR = Path(__file__).resolve().parent

_DEFAULT_ESCUDO_HTTPS = (
    "https://raw.githubusercontent.com/claudiomlarrea/valorador_informes_finales/"
    "main/assets/escudo_uccuyo.png"
)


def _extras_escudo_url() -> str | None:
    """URL HTTPS opcional (Streamlit suele sanitizar data: en Markdown)."""
    v = os.environ.get("UCC_ESCUDO_URL", "").strip()
    if v:
        return v

    try:
        secrets = st.secrets
    except Exception:
        return None

    get = getattr(secrets, "get", None)
    if callable(get):
        for key in ("UCC_ESCUDO_URL", "ucc_escudo_url"):
            raw = get(key)
            if isinstance(raw, str) and raw.strip():
                return raw.strip()

    for key in ("ucc_escudo_url", "UCC_ESCUDO_URL"):
        try:
            raw = secrets[key]
        except Exception:
            continue
        if isinstance(raw, str) and raw.strip():
            return raw.strip()

    return None


def _banner_crest_markup() -> tuple[str, str]:
    """Fragmento CSS + <img HTTPS> del escudo (evita data: … que suele sanitizar Markdown)."""

    extra_css = """
    div.ucci-inst-header .ucci-inst-escudo-img {
        display: block;
        width: auto;
        max-width: 104px;
        max-height: 96px;
        height: auto;
        object-fit: contain;
        background: rgba(255, 255, 255, 0.98);
        border-radius: 10px;
        padding: 0.35rem;
        box-sizing: border-box;
    }
    """
    src = (_extras_escudo_url() or "").strip() or _DEFAULT_ESCUDO_HTTPS
    crest = (
        f'<img class="ucci-inst-escudo-img" src="{src}" '
        f'alt="Universidad Católica de Cuyo" loading="lazy" '
        f'referrerpolicy="no-referrer" />'
    )
    return extra_css, crest


st.set_page_config(layout="wide")

_BANNER_CREST_CSS, _BANNER_CREST_HTML = _banner_crest_markup()

# Llaves "{{" debajo son para interpolar después: sin .replace() Streamlit envía "{{" literal
# y el CSS queda inválido (solo aplica tema config.toml → botón Browse blanco).
_UCCI_STYLE_BLOCK = """
<style>
:root {{
    --ucci-green: #00664d;
    --ucci-green-dark: #00523e;
    --ucci-accent: #28a745;
    --ucci-page-bg: #E6E6E6;
    --ucci-sidebar-bg: #262730;
    --ucci-text: #262730;
}}

.stApp {{
    background-color: var(--ucci-page-bg);
}}

.block-container {{
    padding-top: 1.25rem;
}}

section[data-testid="stSidebar"] {{
    background-color: var(--ucci-sidebar-bg);
}}
[data-testid="stSidebar"] [data-testid="stMarkdown"],
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label {{
    color: rgba(255, 255, 255, 0.92);
}}

/* Banner institucional: caja verde unificada (como Informes de avance) */
.ucci-inst-header {{
    display: flex;
    align-items: center;
    gap: 1.15rem;
    box-sizing: border-box;
    background-color: var(--ucci-green);
    border-radius: 12px;
    padding: 1rem 1.35rem;
    margin-bottom: 1.25rem;
}}

.ucci-inst-banner-text {{
    flex: 1 1 auto;
    min-width: 0;
}}

.ucci-inst-header .ucci-inst-banner-text,
.ucci-inst-header .ucci-inst-banner-text * {{
    color: #ffffff !important;
}}

.ucci-inst-header .ucci-inst-banner-text a {{
    color: #ffffff !important;
    text-decoration: underline dotted rgba(255, 255, 255, 0.55);
}}

.header-ucciuyo h1.ucci-banner-heading,
.header-ucciuyo h2.ucci-banner-heading,
.header-ucciuyo h3.ucci-banner-heading {{
    color: #ffffff !important;
    margin: 0;
    line-height: 1.2;
    font-family: "Source Sans Pro", ui-sans-serif, system-ui, sans-serif;
}}

.header-ucciuyo h1.ucci-banner-heading {{
    font-size: clamp(1.35rem, 2.8vw, 1.95rem);
    font-weight: 700;
}}

.header-ucciuyo h2.ucci-banner-heading {{
    margin-top: 0.5rem !important;
    font-size: clamp(1rem, 2vw, 1.22rem);
    font-weight: 600;
}}

.header-ucciuyo h3.ucci-banner-heading {{
    margin-top: 0.3rem !important;
    font-size: clamp(0.85rem, 1.4vw, 1rem);
    font-weight: 400;
    color: rgba(255, 255, 255, 0.93) !important;
}}

{_BANNER_CREST_CSS}

/* Título de la app tipo tarjeta */
.ucci-main-title-card {{
    background: #ffffff;
    border-radius: 12px;
    padding: 1.1rem 1.35rem 1rem;
    margin-bottom: 1rem;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.07);
    border: 1px solid rgba(0, 38, 28, 0.08);
    box-sizing: border-box;
}}

.ucci-main-title-card h1.ucci-app-title {{
    margin: 0 !important;
    padding: 0 !important;
    font-size: clamp(1.35rem, 2.6vw, 1.72rem);
    font-weight: 700;
    color: var(--ucci-text) !important;
}}

.ucci-main-title-card p.ucci-app-subtitle {{
    margin: 0.55rem 0 0 0 !important;
    font-size: 0.95rem;
    line-height: 1.45;
    color: #5f6368 !important;
}}

/* Contenido general */
h1:not(.ucci-banner-heading):not(.ucci-app-title),
h2:not(.ucci-banner-heading),
h3:not(.ucci-banner-heading),
h4 {{
    color: var(--ucci-green-dark) !important;
}}
p,
label {{
    color: var(--ucci-text) !important;
}}

/* Carga de archivos: franja oscura + CTA verde (alineado a informes de avance + Streamlit ≥1.37) */
[data-testid="stFileUploader"] {{
    background-color: transparent !important;
    border: none !important;
    padding: 0 !important;
    margin-top: 0.15rem !important;
}}
[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] {{
    background-color: #1e1e1e !important;
    border-radius: 12px !important;
    border: 1px solid rgba(255, 255, 255, 0.08) !important;
    padding: 0.85rem 1rem !important;
}}
[data-testid="stBaseButton-primary"],
[data-testid="stBaseButton-secondary"],
[data-testid="stBaseButton-tertiary"] {{
    background-color: var(--ucci-green) !important;
    color: #ffffff !important;
    border-color: transparent !important;
    --text-color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    font-weight: 600 !important;
}}
[data-testid="stBaseButton-primary"]:hover,
[data-testid="stBaseButton-secondary"]:hover,
[data-testid="stBaseButton-tertiary"]:hover {{
    background-color: var(--ucci-green-dark) !important;
    border-color: transparent !important;
    color: #ffffff !important;
    --text-color: #ffffff !important;
}}
[data-testid="stBaseButton-primary"] p,
[data-testid="stBaseButton-primary"] span,
[data-testid="stBaseButton-secondary"] p,
[data-testid="stBaseButton-secondary"] span,
[data-testid="stBaseButton-tertiary"] p,
[data-testid="stBaseButton-tertiary"] span,
[data-testid="stBaseButton-primary"] div,
[data-testid="stBaseButton-secondary"] div,
[data-testid="stBaseButton-tertiary"] div {{
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}}
[data-testid="stBaseButton-primary"] svg,
[data-testid="stBaseButton-secondary"] svg,
[data-testid="stBaseButton-tertiary"] svg,
[data-testid="stFileUploader"] button svg {{
    fill: #ffffff !important;
    color: #ffffff !important;
}}

.stButton > button,
[data-testid="stDownloadButton"] button,
[data-testid="stFileUploader"] button {{
    background-color: var(--ucci-green) !important;
    color: #ffffff !important;
    border-radius: 8px !important;
    border: none !important;
    font-weight: 600 !important;
    --text-color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}}
.stButton > button:hover,
[data-testid="stDownloadButton"] button:hover,
[data-testid="stFileUploader"] button:hover {{
    background-color: var(--ucci-green-dark) !important;
    border-color: transparent !important;
}}
.stButton > button *,
[data-testid="stDownloadButton"] button *,
[data-testid="stFileUploader"] button * {{
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}}

div[data-testid="stAlert"] {{
    border-radius: 10px;
}}

[data-baseweb="slider"] {{
    color: var(--ucci-green);
}}

.stSlider label,
[data-testid="stTextInput"] label,
[data-testid="stTextArea"] label {{
    position: relative;
    padding-left: 1rem;
}}
.stSlider label::before,
[data-testid="stTextInput"] label::before,
[data-testid="stTextArea"] label::before {{
    content: "";
    position: absolute;
    left: 0;
    top: 0.45rem;
    width: 9px;
    height: 9px;
    border-radius: 50%;
    background: var(--ucci-accent);
}}

.stTextInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea {{
    border-radius: 12px !important;
    border: 1px solid rgba(0, 82, 62, 0.22) !important;
    background-color: #ffffff !important;
    color: var(--ucci-text) !important;
    caret-color: var(--ucci-green-dark) !important;
}}

/*
 * Texto claro en la franja oscura sin usar "dropzone *" (evita leyendas blancas fuera del recuadro).
 */
/* Sin "p": si el markdown queda bajo el ancestro del uploader, Dropzone p blanco pisaba la leyenda. */
[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] small {{
    color: rgba(255, 255, 255, 0.92) !important;
    -webkit-text-fill-color: rgba(255, 255, 255, 0.92) !important;
}}

[data-testid="stFileUploader"] label,
[data-testid="stFileUploader"] [data-testid="stWidgetLabel"],
[data-testid="stFileUploader"] [data-testid="stWidgetLabel"] *,
[data-testid="stFileUploader"] label[data-testid="stWidgetLabel"],
[data-testid="stFileUploader"] label[data-testid="stWidgetLabel"] *,
[data-testid="stFileUploader"] [data-baseweb="form-control-label"],
[data-testid="stFileUploader"] [data-baseweb="form-control-label"] * {{
    color: #111111 !important;
    -webkit-text-fill-color: #111111 !important;
}}

[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] button,
[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] button *,
[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] svg {{
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    fill: #ffffff !important;
}}

</style>
"""
st.markdown(
    _UCCI_STYLE_BLOCK.replace("{{", "{")
    .replace("}}", "}")
    .replace("{_BANNER_CREST_CSS}", _BANNER_CREST_CSS),
    unsafe_allow_html=True,
)

# ============================
# CARGA DE RÚBRICA (Anexo VI)
# ============================
with open(_APP_DIR / "rubric_final.yaml", "r", encoding="utf-8") as f:
    config = yaml.safe_load(f)

weights = config["weights"]
thresholds = config["thresholds"]
keywords = config["keywords"]
labels = config.get("labels") or {}


def criterion_label(key: str) -> str:
    """Nombre oficial para tablas y exportación."""
    return labels.get(key, key.replace("_", " ").title())


# ============================
# FUNCIONES
# ============================


def _docx_paragraphs_and_tables(doc: Document) -> str:
    """Anexo III: gran parte del contenido está en tablas."""
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                c = (cell.text or "").strip()
                if c:
                    parts.append(c)
    return "\n".join(parts)


def extract_text(file):
    # Streamlit sólo revisa MIME al subir; el nombre puede llegar como .DOCX / .PDF
    name_lc = (getattr(file, "name", None) or "").lower()
    if name_lc.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text

    if name_lc.endswith(".docx"):
        doc = Document(file)
        return _docx_paragraphs_and_tables(doc)

    return ""


def auto_score(text, keywords_dict):
    """
    Escala 0–4 por criterio según proporción de indicios normativos presentes
    en el texto (rubric_final.yaml). Coherente con ponderación Anexo VI.
    """
    text_low = (text or "").lower()
    scores: dict[str, int] = {}
    for section, keys in keywords_dict.items():
        keylist = [(k or "").strip() for k in keys if (k or "").strip()]
        if not keylist:
            scores[section] = 0
            continue
        hits = sum(1 for k in keylist if k.lower() in text_low)
        ratio = hits / len(keylist)
        if ratio >= 0.45:
            scores[section] = 4
        elif ratio >= 0.30:
            scores[section] = 3
        elif ratio >= 0.15:
            scores[section] = 2
        elif ratio > 0:
            scores[section] = 1
        else:
            scores[section] = 0
    return scores


def weighted_score(scores, weights):
    """Porcentaje según Anexo VI: suma ponderada sobre escala 0–4 (sin descuentos extra)."""
    keys = set(scores) & set(weights)
    if not keys:
        return 0.0
    total = sum(scores[s] * weights[s] for s in keys)
    max_total = sum(weights[s] for s in keys) * 4
    return (total / max_total) * 100 if max_total > 0 else 0.0


def generate_excel(scores, percent, thresholds, label_fn=None):
    label_fn = label_fn or (lambda k: str(k))
    wb = Workbook()
    ws = wb.active
    ws.title = "Resultados"

    ws.append(["Criterio", "Puntaje (0–4)"])
    for k, v in scores.items():
        ws.append([label_fn(k), v])

    ws.append([])
    ws.append(["Puntaje total (%)", round(percent, 2)])

    if percent >= thresholds["aprobado"]:
        result = "Aprobado"
    elif percent >= thresholds["aprobado_obs"]:
        result = "Aprobado con observaciones"
    else:
        result = "No aprobado"

    ws.append(["Dictamen", result])

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def generate_word(scores, percent, thresholds, nombre_proyecto="", label_fn=None):
    label_fn = label_fn or (lambda k: str(k))
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("UCCuyo — Valoración de Informe Final", 1)
    doc.add_paragraph(f"Fecha: {datetime.today().strftime('%Y-%m-%d %H:%M')}")

    if nombre_proyecto:
        doc.add_paragraph(f"Proyecto: {nombre_proyecto}")

    doc.add_heading("Resultados por criterio (Anexo VI)", 2)

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Puntaje (0–4)"

    for k, v in scores.items():
        row = table.add_row().cells
        row[0].text = label_fn(k)
        row[1].text = str(v)

    doc.add_paragraph(f"\nCumplimiento ponderado: {round(percent, 2)}%")

    if percent >= thresholds["aprobado"]:
        result = "Aprobado"
    elif percent >= thresholds["aprobado_obs"]:
        result = "Aprobado con observaciones"
    else:
        result = "No aprobado"

    doc.add_heading("Dictamen", 2)
    doc.add_paragraph(result)

    doc.add_heading("Observaciones del evaluador", 2)
    doc.add_paragraph("." * 78)
    doc.add_paragraph("." * 78)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ============================
# INTERFAZ
# ============================
# Encabezado unificado: franja verde explícita (no depende de :has() del layout de Streamlit).
st.markdown(
    f"""
<div class="ucci-inst-header">
{_BANNER_CREST_HTML}
<div class="ucci-inst-banner-text header-ucciuyo">
<h1 class="ucci-banner-heading">Universidad Católica de Cuyo</h1>
<h2 class="ucci-banner-heading">Secretaría de Investigación</h2>
<h3 class="ucci-banner-heading">Consejo de Investigación</h3>
</div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="ucci-main-title-card">
<h1 class="ucci-app-title">📊 Valorador de Informes Finales</h1>
<p class="ucci-app-subtitle">Subí un informe final (PDF o DOCX) para evaluarlo automáticamente según la rúbrica institucional.</p>
</div>
""",
    unsafe_allow_html=True,
)

st.info("**Cargar archivo** — PDF o DOCX (hasta 200 MB por archivo). Usá el recuadro oscuro abajo.")
archivo = st.file_uploader(
    "Cargar archivo (PDF o DOCX)",
    type=["pdf", "docx"],
    label_visibility="collapsed",
    key="upload_informe_final",
)

if archivo:
    texto = extract_text(archivo)

    with st.expander("Ver texto extraído"):
        st.text_area("Texto completo", texto, height=280)

    st.subheader("Evaluación automática")
    auto_scores = auto_score(texto, keywords)
    ordered_auto = {k: auto_scores[k] for k in weights}

    df = pd.DataFrame(
        [(criterion_label(k), v) for k, v in ordered_auto.items()],
        columns=["Criterio", "Puntaje (0–4)"],
    )
    st.dataframe(df, use_container_width=True)

    auto_percent = weighted_score(ordered_auto, weights)
    st.metric("Puntaje automático inicial (%)", round(auto_percent, 2))
    st.caption(
        "Escala Anexo VI: ≥60 % aprobado; 50–59 % con observaciones; <50 % no aprobado."
    )

    st.subheader("Ajuste manual (opcional)")
    manual_scores: dict[str, int] = {}
    for k in ordered_auto.keys():
        manual_scores[k] = st.slider(
            criterion_label(k),
            0,
            4,
            int(ordered_auto[k]),
            key=f"slider_{k}",
        )

    adjusted_percent = weighted_score(manual_scores, weights)
    st.metric("Puntaje total ajustado (%)", round(adjusted_percent, 2))

    if adjusted_percent >= thresholds["aprobado"]:
        resultado = "✅ Aprobado"
    elif adjusted_percent >= thresholds["aprobado_obs"]:
        resultado = "⚠️ Aprobado con observaciones"
    else:
        resultado = "❌ No aprobado"

    st.success(f"Dictamen (con ajuste manual): {resultado}")

    nombre = st.text_input("Nombre del proyecto (aparecerá en el Word):", "")

    if st.button("Generar informes", type="primary"):
        excel = generate_excel(
            manual_scores, adjusted_percent, thresholds, label_fn=criterion_label
        )
        word = generate_word(
            manual_scores,
            adjusted_percent,
            thresholds,
            nombre,
            label_fn=criterion_label,
        )

        st.download_button(
            "Descargar Excel", excel, "resultado.xlsx", type="primary"
        )
        st.download_button(
            "Descargar Word", word, "resultado.docx", type="primary"
        )

        st.success("Informes generados con los puntajes ajustados.")
